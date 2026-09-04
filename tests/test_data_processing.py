import pandas as pd
import pytest
from docx import Document

from backend.app.data_loader import (
    FileLoadingError,
    load_requirements_excel,
    load_requirements_file,
    load_requirements_word,
    standardize_column_name,
)

from backend.app.validators import (
    DataValidationError,
    validate_requirements_dataframe,
)

from backend.app.text_preprocessor import (
    clean_original_text,
    normalize_text,
    preprocess_text,
)


def test_clean_original_text_removes_extra_whitespace() -> None:
    text = "  Port kontrolü   en fazla 3 kez\ntekrarlanmalıdır.  "

    result = clean_original_text(text)

    assert result == "Port kontrolü en fazla 3 kez tekrarlanmalıdır."


def test_normalize_text_handles_turkish_characters() -> None:
    text = "İŞLEM IŞIĞI KONTROL EDİLMELİDİR."

    result = normalize_text(text)

    assert result == "işlem ışığı kontrol edilmelidir."


def test_normalize_text_preserves_numbers_and_percentages() -> None:
    text = "İŞLEM %80 BAŞARI ORANIYLA 15 DAKİKA İÇİNDE TAMAMLANMALIDIR."

    result = normalize_text(text)

    assert result == (
        "işlem %80 başarı oranıyla 15 dakika içinde tamamlanmalıdır."
    )


def test_preprocess_text_returns_both_versions() -> None:
    text = "  Port kontrolü   EN FAZLA 3 kez\ntekrarlanmalıdır.  "

    result = preprocess_text(text)

    assert result == {
        "original_text": "Port kontrolü EN FAZLA 3 kez tekrarlanmalıdır.",
        "normalized_text": "port kontrolü en fazla 3 kez tekrarlanmalıdır.",
    }


def test_none_value_returns_empty_text() -> None:
    result = preprocess_text(None)

    assert result == {
        "original_text": "",
        "normalized_text": "",
    }

def create_valid_requirements_dataframe() -> pd.DataFrame:
    return pd.DataFrame(
        {
            "requirement_id": ["REQ-001", "REQ-002"],
            "requirement_text": [
                "Port kontrolü yapılmalıdır.",
                "Müşteriye SMS gönderilmelidir.",
            ],
            "module": ["Resource", "Notification"],
            "version": ["1.0", "1.0"],
        }
    )


def test_valid_requirements_dataframe_passes_validation() -> None:
    dataframe = create_valid_requirements_dataframe()

    validate_requirements_dataframe(dataframe)


def test_missing_required_column_raises_error() -> None:
    dataframe = create_valid_requirements_dataframe().drop(columns=["module"])

    with pytest.raises(DataValidationError) as exception:
        validate_requirements_dataframe(dataframe)

    assert "Eksik zorunlu sütunlar: module" in str(exception.value)


def test_blank_requirement_text_raises_error() -> None:
    dataframe = create_valid_requirements_dataframe()
    dataframe.loc[1, "requirement_text"] = "   "

    with pytest.raises(DataValidationError) as exception:
        validate_requirements_dataframe(dataframe)

    assert "1 satırda requirement_text boş." in str(exception.value)


def test_duplicate_requirement_id_raises_error() -> None:
    dataframe = create_valid_requirements_dataframe()
    dataframe.loc[1, "requirement_id"] = "REQ-001"

    with pytest.raises(DataValidationError) as exception:
        validate_requirements_dataframe(dataframe)

    assert "REQ-001" in str(exception.value)


def test_empty_dataframe_raises_error() -> None:
    dataframe = create_valid_requirements_dataframe().iloc[0:0]

    with pytest.raises(DataValidationError) as exception:
        validate_requirements_dataframe(dataframe)

    assert "Dosya en az bir gereksinim içermelidir." in str(exception.value)

def test_standardize_column_name() -> None:
    assert standardize_column_name("Requirement ID") == "requirement_id"
    assert standardize_column_name("Requirement-Text") == "requirement_text"
    assert standardize_column_name(" Module ") == "module"


def test_load_requirements_excel_returns_processed_dataframe(
    tmp_path,
) -> None:
    dataframe = pd.DataFrame(
        {
            "Requirement ID": ["REQ-001"],
            "Requirement Text": [
                "  Port kontrolü   EN FAZLA 3 kez yapılmalıdır. "
            ],
            "Module": ["Resource"],
            "Version": ["1.0"],
        }
    )

    file_path = tmp_path / "requirements.xlsx"
    dataframe.to_excel(file_path, index=False)

    result = load_requirements_excel(file_path)

    assert result.loc[0, "requirement_id"] == "REQ-001"
    assert result.loc[0, "original_text"] == (
        "Port kontrolü EN FAZLA 3 kez yapılmalıdır."
    )
    assert result.loc[0, "normalized_text"] == (
        "port kontrolü en fazla 3 kez yapılmalıdır."
    )


def test_load_requirements_excel_missing_file_raises_error(
    tmp_path,
) -> None:
    file_path = tmp_path / "missing.xlsx"

    with pytest.raises(FileLoadingError) as exception:
        load_requirements_excel(file_path)

    assert "Dosya bulunamadı" in str(exception.value)


def test_load_requirements_excel_rejects_unsupported_extension(
    tmp_path,
) -> None:
    file_path = tmp_path / "requirements.csv"
    file_path.write_text("test", encoding="utf-8")

    with pytest.raises(FileLoadingError) as exception:
        load_requirements_excel(file_path)

    assert "Yalnızca .xlsx" in str(exception.value)


def test_load_requirements_excel_validates_required_columns(
    tmp_path,
) -> None:
    dataframe = pd.DataFrame(
        {
            "Requirement ID": ["REQ-001"],
            "Requirement Text": ["Port kontrolü yapılmalıdır."],
        }
    )

    file_path = tmp_path / "requirements.xlsx"
    dataframe.to_excel(file_path, index=False)

    with pytest.raises(DataValidationError) as exception:
        load_requirements_excel(file_path)

    assert "Eksik zorunlu sütunlar" in str(exception.value)


def test_load_requirements_word_table_returns_processed_dataframe(
    tmp_path,
) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=4)
    for index, value in enumerate(
        ["Requirement ID", "Requirement", "Module", "Version"]
    ):
        table.rows[0].cells[index].text = value

    row = table.add_row().cells
    for index, value in enumerate(
        ["REQ-001", "The user can sign in", "Authentication", "1.0"]
    ):
        row[index].text = value

    file_path = tmp_path / "requirements.docx"
    document.save(file_path)

    result = load_requirements_word(file_path)

    assert result.loc[0, "requirement_id"] == "REQ-001"
    assert result.loc[0, "requirement_text"] == "The user can sign in"
    assert result.loc[0, "module"] == "Authentication"


def test_load_requirements_word_accepts_preview_mapping(
    tmp_path,
) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=4)
    for index, value in enumerate(
        ["Gereksinim ID", "Gereksinim", "Modül", "Versiyon"]
    ):
        table.rows[0].cells[index].text = value
    row = table.add_row().cells
    for index, value in enumerate(
        ["REQ-001", "Kullanıcı giriş yapabilmelidir.", "Kimlik", "1.0"]
    ):
        row[index].text = value
    file_path = tmp_path / "requirements.docx"
    document.save(file_path)

    result = load_requirements_word(
        file_path,
        column_mapping={
            "requirement_id": "requirement_id",
            "requirement_text": "requirement_text",
            "module": "module",
            "version": "version",
        },
    )

    assert result.loc[0, "requirement_text"] == "Kullanıcı giriş yapabilmelidir."


def test_load_requirements_file_reads_word_paragraphs(
    tmp_path,
) -> None:
    document = Document()
    document.add_paragraph("The user can sign in")
    document.add_paragraph("The user can reset a password")
    file_path = tmp_path / "requirements.docx"
    document.save(file_path)

    result = load_requirements_file(file_path)

    assert len(result.index) == 2
    assert result["requirement_id"].tolist() == ["AUTO-0001", "AUTO-0002"]
    assert result["module"].tolist() == ["General", "General"]


def test_load_requirements_word_supports_headerless_one_column_table(
    tmp_path,
) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Kullanıcı giriş yapabilmelidir."
    table.cell(1, 0).text = "Kullanıcı çıkış yapabilmelidir."
    file_path = tmp_path / "requirements.docx"
    document.save(file_path)

    result = load_requirements_word(file_path)

    assert result["requirement_text"].tolist() == [
        "Kullanıcı giriş yapabilmelidir.",
        "Kullanıcı çıkış yapabilmelidir.",
    ]
