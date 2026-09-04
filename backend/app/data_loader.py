import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml.ns import qn

from backend.app.text_preprocessor import preprocess_text
from backend.app.validators import validate_requirements_dataframe


class FileLoadingError(ValueError):
    """Gereksinim dosyası okunamadığında oluşan hatayı temsil eder."""


SUPPORTED_REQUIREMENT_EXTENSIONS = {".xlsx", ".docx"}


REQUIREMENT_COLUMN_ALIASES = {
    "requirement_id": {
        "requirementid", "reqid", "id", "kod", "gereksinimid", "gereksinimno", "ref", "key",
    },
    "requirement_text": {
        "requirementtext", "requirement", "requirementdescription", "description", "details", "text", "subject", "title", "aciklama", "gereksinim", "gereksinimmetni", "gereksinimaciklamasi", "gereksinimicerigi", "madde", "isregeli", "isgereksinimi",
    },
    "module": {
        "module", "component", "area", "category", "modul", "menu", "menukategori", "fonksiyon",
    },
    "version": {
        "version", "release", "sprint", "versiyon", "surum", "surumno",
    },
}


def _header_key(value: object) -> str:
    text = str(value).strip().lower()
    replacements = str.maketrans("ıişğüöç", "iisguoc")
    text = text.translate(replacements)
    return re.sub(r"[^a-z0-9]+", "", text)


def detect_requirement_columns(columns: list[object]) -> dict[str, str | None]:
    """Farklı şirket kolon adlarını ScopeDiff alanlarına eşler."""
    normalized = {_header_key(column): str(column) for column in columns}
    mapping: dict[str, str | None] = {}

    for canonical, aliases in REQUIREMENT_COLUMN_ALIASES.items():
        mapping[canonical] = next(
            (normalized[alias] for alias in aliases if alias in normalized),
            None,
        )

    return mapping


def standardize_column_name(column_name: object) -> str:
    """
    Sütun adlarını standart formata dönüştürür.

    Örnek:
    'Requirement ID' -> 'requirement_id'
    'Requirement-Text' -> 'requirement_text'
    """
    name = str(column_name).strip().lower()
    name = name.translate(str.maketrans("ıişğüöç", "iisguoc"))
    name = re.sub(r"[\s\-]+", "_", name)
    name = re.sub(r"[^a-z0-9_]", "", name)
    name = re.sub(r"_+", "_", name)

    return name.strip("_")


def standardize_column_names(dataframe: pd.DataFrame) -> pd.DataFrame:
    """DataFrame sütunlarının standartlaştırılmış bir kopyasını döndürür."""
    result = dataframe.copy()
    result.columns = [
        standardize_column_name(column)
        for column in result.columns
    ]

    return result


def load_requirements_excel(
    file_path: str | Path,
    sheet_name: str | int = 0,
    column_mapping: dict[str, str] | None = None,
    version_fallback: str | None = None,
) -> pd.DataFrame:
    """
    Gereksinim Excel dosyasını okur, doğrular ve metinleri işler.

    Dönen tabloda şu ek sütunlar bulunur:

    - original_text
    - normalized_text
    """
    path = Path(file_path)

    if not path.exists():
        raise FileLoadingError(f"Dosya bulunamadı: {path}")

    if path.suffix.lower() != ".xlsx":
        raise FileLoadingError(
            "Yalnızca .xlsx uzantılı Excel dosyaları desteklenmektedir."
        )

    try:
        dataframe = pd.read_excel(
            path,
            sheet_name=sheet_name,
            engine="openpyxl",
        )
    except ValueError as error:
        raise FileLoadingError(
            f"Excel çalışma sayfası okunamadı: {error}"
        ) from error
    except Exception as error:
        raise FileLoadingError(
            f"Excel dosyası okunamadı: {error}"
        ) from error

    dataframe = standardize_column_names(dataframe)

    detected_mapping = detect_requirement_columns(list(dataframe.columns))
    selected_mapping = {
        **detected_mapping,
        **(column_mapping or {}),
    }

    text_column = selected_mapping.get("requirement_text")
    if not text_column or text_column not in dataframe.columns:
        raise FileLoadingError(
            "Gereksinim metni kolonu bulunamadı. "
            "Description, Requirement veya Açıklama benzeri bir kolon seçin."
        )

    rename_map = {
        source: canonical
        for canonical, source in selected_mapping.items()
        if source and source in dataframe.columns
    }
    dataframe = dataframe.rename(columns=rename_map)

    row_count = len(dataframe.index)
    if "requirement_id" not in dataframe.columns:
        dataframe["requirement_id"] = ""
    dataframe["requirement_id"] = dataframe["requirement_id"].fillna("")
    blank_ids = dataframe["requirement_id"].astype(str).str.strip().eq("")
    dataframe.loc[blank_ids, "requirement_id"] = [
        f"AUTO-{index:04d}"
        for index in range(1, int(blank_ids.sum()) + 1)
    ]
    validate_requirements_dataframe(dataframe)

    processed_texts = dataframe["requirement_text"].apply(preprocess_text)

    dataframe["original_text"] = processed_texts.apply(
        lambda item: item["original_text"]
    )
    dataframe["normalized_text"] = processed_texts.apply(
        lambda item: item["normalized_text"]
    )

    dataframe["requirement_id"] = (
        dataframe["requirement_id"]
        .astype(str)
        .str.strip()
    )

    dataframe["module"] = (
        dataframe["module"]
        .astype(str)
        .str.strip()
    )

    dataframe["version"] = (
        dataframe["version"]
        .astype(str)
        .str.strip()
    )

    return dataframe


def _finalize_word_dataframe(
    dataframe: pd.DataFrame,
    version_fallback: str | None = None,
    column_mapping: dict[str, str] | None = None,
) -> pd.DataFrame:
    """Word'den gelen tabloyu analiz pipeline'ının ortak formatına çevirir."""
    if dataframe.empty:
        raise FileLoadingError("Word dosyası en az bir gereksinim içermelidir.")

    dataframe = standardize_column_names(dataframe)
    detected_mapping = detect_requirement_columns(list(dataframe.columns))
    selected_mapping = dict(detected_mapping)
    for canonical, source in (column_mapping or {}).items():
        if not source:
            continue
        normalized_source = standardize_column_name(source)
        if normalized_source in dataframe.columns:
            selected_mapping[canonical] = normalized_source

    text_column = selected_mapping.get("requirement_text")
    if not text_column or text_column not in dataframe.columns:
        raise FileLoadingError(
            "Word dosyasında gereksinim metni bulunamadı. "
            "Gereksinimleri paragraf olarak veya bir tabloda 'Requirement' "
            "/ 'Description' başlığı altında ekleyin."
        )

    dataframe = dataframe.rename(
        columns={
            source: canonical
            for canonical, source in selected_mapping.items()
            if source and source in dataframe.columns
        }
    )

    if "requirement_id" not in dataframe.columns:
        dataframe["requirement_id"] = ""
    if "module" not in dataframe.columns:
        dataframe["module"] = "General"
    if "version" not in dataframe.columns:
        dataframe["version"] = version_fallback or "document"

    dataframe["requirement_id"] = dataframe["requirement_id"].fillna("")
    blank_ids = dataframe["requirement_id"].astype(str).str.strip().eq("")
    dataframe.loc[blank_ids, "requirement_id"] = [
        f"AUTO-{index:04d}"
        for index in range(1, int(blank_ids.sum()) + 1)
    ]

    validate_requirements_dataframe(dataframe)
    processed_texts = dataframe["requirement_text"].apply(preprocess_text)
    dataframe["original_text"] = processed_texts.apply(
        lambda item: item["original_text"]
    )
    dataframe["normalized_text"] = processed_texts.apply(
        lambda item: item["normalized_text"]
    )

    for column in ("requirement_id", "module", "version"):
        dataframe[column] = dataframe[column].astype(str).str.strip()

    return dataframe


def load_requirements_word(
    file_path: str | Path,
    column_mapping: dict[str, str] | None = None,
    version_fallback: str | None = None,
) -> pd.DataFrame:
    """Word paragraflarını veya tablolarını gereksinim DataFrame'ine çevirir."""
    path = Path(file_path)
    if not path.exists():
        raise FileLoadingError(f"Dosya bulunamadı: {path}")
    if path.suffix.lower() != ".docx":
        raise FileLoadingError("Yalnızca .docx uzantılı Word dosyaları desteklenmektedir.")

    try:
        document = Document(path)
        table_rows: list[list[str]] = []
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    table_rows.append(values)

        if table_rows:
            width = max(len(row) for row in table_rows)
            rows = [row + [""] * (width - len(row)) for row in table_rows]
            header_mapping = detect_requirement_columns(rows[0])
            has_header = bool(header_mapping.get("requirement_text"))
            if has_header:
                columns = [value or f"Column {index + 1}" for index, value in enumerate(rows[0])]
                dataframe = pd.DataFrame(rows[1:], columns=columns)
            else:
                if width == 1:
                    columns = ["requirement_text"]
                elif width == 2:
                    columns = ["requirement_id", "requirement_text"]
                else:
                    columns = [
                        "requirement_id",
                        "requirement_text",
                        "module",
                        "version",
                    ][:width]
                columns += [f"Column {index + 1}" for index in range(len(columns), width)]
                dataframe = pd.DataFrame(rows, columns=columns)
        else:
            paragraphs = [
                paragraph.text.strip()
                for paragraph in document.paragraphs
                if paragraph.text.strip()
            ]
            if not paragraphs:
                text_nodes = [
                    node.text.strip()
                    for node in document.element.body.iter(qn("w:t"))
                    if node.text and node.text.strip()
                ]
                paragraphs = text_nodes
            dataframe = pd.DataFrame({"requirement_text": paragraphs})

        return _finalize_word_dataframe(
            dataframe,
            version_fallback=version_fallback,
            column_mapping=column_mapping,
        )
    except FileLoadingError:
        raise
    except Exception as error:
        raise FileLoadingError(f"Word dosyası okunamadı: {error}") from error


def load_requirements_file(
    file_path: str | Path,
    sheet_name: str | int = 0,
    column_mapping: dict[str, str] | None = None,
    version_fallback: str | None = None,
) -> pd.DataFrame:
    """Excel veya Word gereksinim dosyasını ortak analiz formatında yükler."""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".xlsx":
        return load_requirements_excel(
            file_path,
            sheet_name=sheet_name,
            column_mapping=column_mapping,
            version_fallback=version_fallback,
        )
    if suffix == ".docx":
        return load_requirements_word(
            file_path,
            column_mapping=column_mapping,
            version_fallback=version_fallback,
        )
    raise FileLoadingError("Yalnızca .xlsx Excel veya .docx Word dosyaları desteklenmektedir.")
